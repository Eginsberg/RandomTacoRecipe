import requests
import docx
from PIL import Image, ImageFont, ImageDraw

# I start by importing the needed libraries

image = Image.open('taco.jpg')

width = image.width
height = image.height

'''width = image.width

height = image.height

print(width, height)

I then use the commented out code to find the proportions of the image to change by a ratio of 7.3225 to get the
max size to 800 px'''

scaled_width = int(width / 7.3225)
scaled_height = int(height / 7.3225)

# I then scaled the size down to meet the standards

image = image.resize((scaled_width, scaled_height))

# remade the image

img_draw = ImageDraw.Draw(image)
font = ImageFont.truetype('arial.ttf', 50)
img_draw.text([120, 20], 'Random Taco Cookbook', fill='black', font=font)

# I then wrote the words random taco cookbook on the image

image.save('new_taco.jpg')

# I then saved the image

data = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()

# I start the next part by requesting the data

document = docx.Document()

# I then create the initial document

header = document.sections[0].header
header.add_paragraph('Random Taco Cookbook')

# Adding the header

document.add_picture('new_taco.jpg')

# Adding the picture

document.add_paragraph('Image author : Tai\"s Captures')

# Giving the image author credit

document.add_paragraph('https://taco-1150.herokuapp.com/random/?full_taco=true')

# Adding random taco API link

document.add_paragraph('Project made by Ethan Ginsberg')

# Adding my own name

document.add_page_break()

# Inserting page break

document.save('Taco_Recipe_Book.docx')

# I make sure to actually save the document

